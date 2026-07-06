import csv
import io
import os
import re
import zipfile
from typing import Any, Dict, List

from app.attachment_llm_summary import summarize_attachment_with_llm

RISKY_EXTENSIONS = {
    ".exe", ".bat", ".cmd", ".scr", ".js", ".vbs", ".ps1", ".msi", ".jar", ".com"
}
ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z", ".tar", ".gz"}

AMOUNT_RE = re.compile(
    r"(?i)(?:\$|usd\s*)\s?\d{1,3}(?:,\d{3})*(?:\.\d{2})?|\b\d+(?:\.\d{2})?\s?(?:usd|dollars)\b"
)
DATE_RE = re.compile(
    r"(?i)\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{1,2},?\s+\d{2,4}\b"
    r"|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
    r"|\b(?:today|tomorrow|tonight|asap|urgent|due date|deadline|start date|effective date)\b"
)
ACTION_RE = re.compile(
    r"(?i)\b(pay|sign|review|approve|submit|send|share|upload|complete|confirm|verify|schedule|reply|respond|accept|return)\b"
)

DOC_PATTERNS = [
    ("certificate", ["certificate", "certification", "course completed", "completed by", "pdu", "contact hours", "certificate id", "registered education provider"]),
    # Travel is intentionally strict. A single word like confirmation/release must not become Travel Document.
    ("travel_document", ["flight", "airline", "airport", "boarding", "boarding pass", "itinerary", "reservation", "booking reference", "pnr", "departure", "arrival", "gate", "seat", "check-in", "trip confirmation"]),
    ("project_document", ["development", "release", "feature", "bug", "frontend", "backend", "fastapi", "react", "ai email", "agent", "implementation", "roadmap", "architecture", "api", "redis", "celery", "ollama", "llm", "rag", "ocr"]),
    ("offer_letter", ["offer letter", "employment offer", "start date", "salary", "compensation", "accept this offer"]),
    ("contract", ["agreement", "contract", "terms and conditions", "effective date", "party", "parties", "signature"]),
    ("invoice", ["invoice", "amount due", "payment due", "bill to", "invoice number", "subtotal", "total due"]),
    ("resume", ["resume", "curriculum vitae", "work experience", "education", "skills", "projects", "linkedin"]),
    ("tax_document", ["tax", "w-2", "1099", "form 1040", "irs", "taxpayer", "withholding"]),
    ("bank_statement", ["bank statement", "account number", "statement period", "available balance", "transaction"]),
    ("medical_document", ["patient", "doctor", "clinic", "hospital", "diagnosis", "prescription", "medical"]),
]

# Speed-first defaults. These keep the clicked Analyze Attachment path fast.
# Increase these only for full/background document analysis.
ATTACHMENT_TEXT_PAGES = int(os.getenv("ATTACHMENT_TEXT_PAGES", "6"))
ATTACHMENT_OCR_PAGES = int(os.getenv("ATTACHMENT_OCR_PAGES", "2"))
ATTACHMENT_DOCX_PARAGRAPHS = int(os.getenv("ATTACHMENT_DOCX_PARAGRAPHS", "80"))
ATTACHMENT_TABLE_ROWS = int(os.getenv("ATTACHMENT_TABLE_ROWS", "12"))
ATTACHMENT_EXTRACTED_TEXT_LIMIT = int(os.getenv("ATTACHMENT_EXTRACTED_TEXT_LIMIT", "12000"))
ATTACHMENT_LLM_TEXT_LIMIT = int(os.getenv("ATTACHMENT_LLM_TEXT_LIMIT", "6000"))

DOC_LABELS = {
    "certificate": "Certificate",
    "travel_document": "Travel Document",
    "offer_letter": "Offer Letter",
    "contract": "Contract",
    "invoice": "Invoice",
    "resume": "Resume",
    "tax_document": "Tax Document",
    "bank_statement": "Bank Statement",
    "medical_document": "Medical Document",
    "general_document": "General Document",
    "project_document": "Project Document",
    "technical_document": "Technical Document",
    "spreadsheet": "Spreadsheet",
    "image": "Image",
    "screenshot": "Screenshot",
    "archive": "Archive",
    "video": "Video",
    "risky_executable": "Risky File",
}


def _ext(filename: str) -> str:
    name = (filename or "").lower().strip()
    return "" if "." not in name else "." + name.rsplit(".", 1)[-1]


def classify_attachment(filename: str, mime_type: str = "") -> str:
    fn = (filename or "").lower()
    mt = (mime_type or "").lower()
    ext = _ext(fn)
    if ext in RISKY_EXTENSIONS:
        return "risky_executable"
    if ext in ARCHIVE_EXTENSIONS or "zip" in mt:
        return "archive"
    if "pdf" in mt or ext == ".pdf":
        return "pdf"
    if "word" in mt or "officedocument.wordprocessingml" in mt or ext in {".doc", ".docx"}:
        return "word"
    if "spreadsheet" in mt or "excel" in mt or ext in {".xls", ".xlsx"}:
        return "excel"
    if "csv" in mt or ext == ".csv":
        return "csv"
    if mt.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}:
        return "image"
    if mt.startswith("video/") or ext in {".mp4", ".mov", ".avi", ".mkv"}:
        return "video"
    if mt.startswith("text/") or ext in {".txt", ".md"}:
        return "text"
    return "other"


def _safe_text(s: str, limit: int = ATTACHMENT_EXTRACTED_TEXT_LIMIT) -> str:
    s = (s or "").replace("\x00", " ")
    s = re.sub(r"[ \t\r\f\v]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()[:limit]


def _configure_tesseract() -> None:
    """Use TESSERACT_CMD from .env/OS when Windows does not expose tesseract in PATH."""
    try:
        import os
        import pytesseract
        cmd = os.getenv("TESSERACT_CMD") or os.getenv("PYTESSERACT_CMD")
        if cmd:
            pytesseract.pytesseract.tesseract_cmd = cmd
    except Exception:
        pass


def _ocr_image_obj(img: Any) -> str:
    try:
        _configure_tesseract()
        import pytesseract
        return pytesseract.image_to_string(img) or ""
    except Exception:
        return ""


def _render_pdf_with_pymupdf(data: bytes, max_pages: int = ATTACHMENT_OCR_PAGES) -> List[Any]:
    """Render PDF pages without Poppler. Works if PyMuPDF is installed."""
    images: List[Any] = []
    try:
        import fitz  # PyMuPDF
        from PIL import Image

        doc = fitz.open(stream=data, filetype="pdf")
        for page_index in range(min(max_pages, len(doc))):
            page = doc.load_page(page_index)
            matrix = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        doc.close()
    except Exception:
        pass
    return images


def _render_pdf_with_pypdfium2(data: bytes, max_pages: int = ATTACHMENT_OCR_PAGES) -> List[Any]:
    """Second no-Poppler renderer. Works if pypdfium2 is installed."""
    images: List[Any] = []
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(data)
        count = min(max_pages, len(pdf))
        for page_index in range(count):
            page = pdf[page_index]
            bitmap = page.render(scale=2.0)
            images.append(bitmap.to_pil())
            page.close()
        pdf.close()
    except Exception:
        pass
    return images


def _render_pdf_with_pdf2image(data: bytes, max_pages: int = ATTACHMENT_OCR_PAGES) -> List[Any]:
    """Poppler-based renderer. Uses POPPLER_PATH from .env/OS on Windows if provided."""
    try:
        import os
        from pdf2image import convert_from_bytes
        poppler_path = os.getenv("POPPLER_PATH") or None
        return convert_from_bytes(
            data,
            first_page=1,
            last_page=max_pages,
            dpi=220,
            poppler_path=poppler_path,
        )
    except Exception:
        return []


def _ocr_pdf_pages(data: bytes, max_pages: int = ATTACHMENT_OCR_PAGES) -> str:
    """
    OCR fallback for scanned/image PDFs.
    Tries PyMuPDF first because it does not need Poppler, then pypdfium2, then pdf2image/Poppler.
    OCR still requires Tesseract installed and reachable through PATH or TESSERACT_CMD.
    """
    images = _render_pdf_with_pymupdf(data, max_pages=max_pages)
    if not images:
        images = _render_pdf_with_pypdfium2(data, max_pages=max_pages)
    if not images:
        images = _render_pdf_with_pdf2image(data, max_pages=max_pages)

    chunks: List[str] = []
    for img in images:
        text = _ocr_image_obj(img)
        if text.strip():
            chunks.append(text)
    return _safe_text("\n".join(chunks))


def _extract_pdf(data: bytes) -> str:
    chunks: List[str] = []

    # 1) Normal embedded-text PDFs
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        for page in reader.pages[:ATTACHMENT_TEXT_PAGES]:
            chunks.append(page.extract_text() or "")
    except Exception:
        pass

    if not "\n".join(chunks).strip():
        try:
            from PyPDF2 import PdfReader as PdfReader2
            reader = PdfReader2(io.BytesIO(data))
            for page in reader.pages[:ATTACHMENT_TEXT_PAGES]:
                chunks.append(page.extract_text() or "")
        except Exception:
            pass

    if not "\n".join(chunks).strip():
        try:
            from pdfminer.high_level import extract_text
            chunks.append(extract_text(io.BytesIO(data), maxpages=ATTACHMENT_TEXT_PAGES) or "")
        except Exception:
            pass

    text = _safe_text("\n".join(chunks))
    if text.strip():
        return text

    # 2) Scanned/image PDF OCR fallback
    return _ocr_pdf_pages(data, max_pages=ATTACHMENT_OCR_PAGES)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        chunks = []
        for p in doc.paragraphs[:ATTACHMENT_DOCX_PARAGRAPHS]:
            if p.text:
                chunks.append(p.text)
        for table in doc.tables[:3]:
            for row in table.rows[:ATTACHMENT_TABLE_ROWS]:
                chunks.append(" | ".join(cell.text for cell in row.cells[:8]))
        return _safe_text("\n".join(chunks))
    except Exception:
        return ""


def _extract_csv(data: bytes) -> str:
    try:
        text = data.decode("utf-8", errors="ignore")
        rows = list(csv.reader(io.StringIO(text)))
        preview = rows[:ATTACHMENT_TABLE_ROWS]
        return _safe_text("\n".join(" | ".join(row[:10]) for row in preview))
    except Exception:
        return ""


def _extract_excel(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        chunks = []
        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            chunks.append(f"Sheet: {sheet_name}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= ATTACHMENT_TABLE_ROWS:
                    break
                chunks.append(" | ".join("" if v is None else str(v) for v in row[:10]))
        return _safe_text("\n".join(chunks))
    except Exception:
        return ""


def _extract_image(data: bytes) -> str:
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        meta = f"Image detected. Format: {img.format}. Size: {img.width}x{img.height}."
        ocr = _ocr_image_obj(img)
        if ocr.strip():
            return _safe_text(meta + "\nOCR text:\n" + ocr)
        return meta + " OCR text was not available. Install Tesseract OCR or set TESSERACT_CMD."
    except Exception:
        return "Image attachment detected. OCR text was not available. Install Tesseract OCR or set TESSERACT_CMD."


def _extract_archive(data: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            names = z.namelist()[:30]
            return "Archive contains: " + ", ".join(names)
    except Exception:
        return "Archive attachment detected. Preview skipped for safety."


def extract_attachment_text(filename: str, mime_type: str, data: bytes) -> str:
    file_type = classify_attachment(filename, mime_type)
    if file_type == "pdf":
        return _extract_pdf(data)
    if file_type == "word":
        return _extract_docx(data)
    if file_type == "csv":
        return _extract_csv(data)
    if file_type == "excel":
        return _extract_excel(data)
    if file_type == "image":
        return _extract_image(data)
    if file_type == "archive":
        return _extract_archive(data)
    if file_type == "text":
        return _safe_text(data.decode("utf-8", errors="ignore"))
    return ""


def _find_amounts(text: str) -> List[str]:
    return list(dict.fromkeys(AMOUNT_RE.findall(text or "")))[:10]


def _find_dates(text: str) -> List[str]:
    return list(dict.fromkeys(DATE_RE.findall(text or "")))[:10]


def _find_actions(text: str) -> List[str]:
    low = (text or "").lower()
    actions: List[str] = []
    for verb in ACTION_RE.findall(low):
        actions.append(f"Possible action: {verb}")
    if "invoice" in low or "payment" in low or "amount due" in low:
        actions.append("Review payment or invoice details.")
    if "sign" in low or "signature" in low or "accept this offer" in low:
        actions.append("Check if signature or acceptance is required.")
    if "deadline" in low or "due date" in low or "start date" in low:
        actions.append("Check important date or deadline.")
    if "resume" in low or "curriculum vitae" in low:
        actions.append("Review candidate profile or resume details.")
    return list(dict.fromkeys(actions))[:8]


def _meaningful_image_text(text: str) -> bool:
    cleaned = re.sub(r"Image detected\. Format:.*?(?:OCR text:)?", " ", text or "", flags=re.I | re.S)
    words = re.findall(r"[A-Za-z0-9]{3,}", cleaned)
    meta_words = {"image", "detected", "format", "size", "ocr", "text", "available", "install", "tesseract"}
    useful = [w for w in words if w.lower() not in meta_words]
    return len(useful) >= int(os.getenv("IMAGE_USEFUL_OCR_WORDS", "6"))


def detect_document_type(filename: str, file_type: str, text: str) -> Dict[str, Any]:
    low = f"{filename or ''}\n{text or ''}".lower()
    if file_type in {"excel", "csv"}:
        return {"document_type": "spreadsheet", "document_label": DOC_LABELS["spreadsheet"], "confidence": 0.80, "matched_terms": []}
    if file_type == "image":
        image_type = "screenshot" if _meaningful_image_text(text) else "image"
        return {"document_type": image_type, "document_label": DOC_LABELS[image_type], "confidence": 0.62 if image_type == "screenshot" else 0.45, "matched_terms": []}
    if file_type in {"archive", "video", "risky_executable"}:
        return {"document_type": file_type, "document_label": DOC_LABELS.get(file_type, file_type.title()), "confidence": 0.80, "matched_terms": []}

    scores: List[tuple] = []
    for doc_type, terms in DOC_PATTERNS:
        hits = [t for t in terms if t in low]
        # Travel needs multiple strong signals. Do not classify as travel from one weak term.
        if doc_type == "travel_document" and len(hits) < 2:
            continue
        if hits:
            scores.append((len(hits), doc_type, hits))

    if not scores:
        return {"document_type": "general_document", "document_label": DOC_LABELS["general_document"], "confidence": 0.35, "matched_terms": []}

    scores.sort(reverse=True, key=lambda x: x[0])
    hit_count, best_type, best_terms = scores[0]
    confidence = min(0.92, 0.42 + (0.10 * hit_count))
    if best_type == "travel_document":
        confidence = max(confidence, 0.68)
    if best_type == "project_document":
        confidence = max(confidence, 0.62)
    return {
        "document_type": best_type,
        "document_label": DOC_LABELS.get(best_type, "General Document"),
        "confidence": round(confidence, 2),
        "matched_terms": best_terms[:8],
    }


def attachment_risk(filename: str, mime_type: str, sender_band: str = "", source_folder: str = "", document_type: str = "") -> Dict[str, Any]:
    file_type = classify_attachment(filename, mime_type)
    band = (sender_band or "").upper()
    folder = (source_folder or "").lower()
    doc_type = (document_type or "").lower()
    reasons: List[str] = []
    score = 0.05

    if file_type == "risky_executable":
        score = 0.95
        reasons.append("Executable or script attachment is risky.")
    elif file_type == "archive":
        score = 0.55
        reasons.append("Archive file should be opened carefully.")
    elif file_type in {"pdf", "word", "excel", "csv"}:
        score = 0.16

    if doc_type in {"invoice", "tax_document", "bank_statement"}:
        score += 0.12
        reasons.append("Financial or sensitive document type detected.")
    if band in {"UNKNOWN", "BULK", "PLATFORM"}:
        score += 0.20
        reasons.append("Sender is not strongly trusted.")
    if folder == "spam":
        score += 0.25
        reasons.append("Attachment came from Spam folder.")

    score = min(score, 1.0)
    level = "high" if score >= 0.70 else "medium" if score >= 0.35 else "low"
    return {"risk_score": round(score, 2), "risk_level": level, "risk_reasons": reasons}


def attachment_priority_boost(document_type: str, risk_level: str, action_items: List[str], confidence: float = 0.60) -> float:
    doc = (document_type or "").lower()
    confidence = max(0.0, min(1.0, float(confidence or 0.0)))
    boost = 0.0
    if doc == "offer_letter":
        boost += 0.25
    elif doc == "contract":
        boost += 0.18
    elif doc == "invoice":
        boost += 0.16
    elif doc in {"tax_document", "bank_statement", "medical_document"}:
        boost += 0.14
    elif doc == "certificate":
        boost += 0.08
    elif doc == "travel_document":
        boost += 0.10
    elif doc == "resume":
        boost += 0.06
    elif doc in {"project_document", "technical_document"}:
        boost += 0.05
    elif doc in {"image", "screenshot", "general_document"}:
        boost += 0.0
    meaningful_action = any(
        ("no action" not in str(x).lower()) and ("not required" not in str(x).lower()) and ("open the image manually" not in str(x).lower())
        for x in (action_items or [])
    )
    if meaningful_action:
        boost += 0.05
    if risk_level == "high":
        boost += 0.10
    elif risk_level == "medium":
        boost += 0.04
    # Low-confidence classifications should not inflate priority.
    if confidence < 0.55:
        boost *= 0.35
    elif confidence < 0.70:
        boost *= 0.70
    return round(min(boost, 0.30), 2)


def _summary_from_fields(file_type: str, doc: Dict[str, Any], extracted: str, dates: List[str], amounts: List[str], actions: List[str]) -> str:
    label = doc.get("document_label") or "Document"
    if extracted:
        first = extracted[:650] + ("..." if len(extracted) > 650 else "")
        prefix = f"{label} detected. "
        if dates:
            prefix += f"Dates found: {', '.join(dates[:3])}. "
        if amounts:
            prefix += f"Amounts found: {', '.join(amounts[:3])}. "
        if actions:
            prefix += f"Action likely needed. "
        return prefix + first
    return (
        f"{label} attachment detected, but readable text was not available. "
        "If this is a scanned PDF/image, install Tesseract OCR and one PDF renderer: PyMuPDF or pypdfium2. "
        "On Windows, set TESSERACT_CMD if Tesseract is not in PATH."
    )


def analyze_attachment_bytes(
    filename: str,
    mime_type: str,
    data: bytes,
    sender_band: str = "",
    source_folder: str = "",
    email_subject: str = "",
    email_sender: str = "",
    email_snippet: str = "",
) -> Dict[str, Any]:
    file_type = classify_attachment(filename, mime_type)

    if file_type == "risky_executable":
        risk = attachment_risk(filename, mime_type, sender_band, source_folder, "risky_executable")
        return {
            "filename": filename,
            "file_type": file_type,
            "document_type": "risky_executable",
            "document_label": "Risky File",
            "document_confidence": 0.95,
            "summary": "Unsafe executable/script attachment. Analysis skipped.",
            "extracted_text": "",
            "action_items": [],
            "dates": [],
            "amounts": [],
            "priority_boost": 0.10,
            "reply_context": "The attachment appears unsafe, so avoid opening it unless verified.",
            **risk,
            "safe_to_preview": False,
        }

    extracted = extract_attachment_text(filename, mime_type, data)
    doc = detect_document_type(filename, file_type, extracted)
    dates = _find_dates(extracted)
    amounts = _find_amounts(extracted)
    actions = _find_actions(extracted)
    llm = summarize_attachment_with_llm(
        filename=filename,
        file_type=file_type,
        detected_document_label=doc.get("document_label", "General Document"),
        extracted_text=extracted[:ATTACHMENT_LLM_TEXT_LIMIT],
        email_subject=email_subject,
        email_sender=email_sender,
        email_snippet=email_snippet,
        dates=dates,
        amounts=amounts,
    )

    # Let the LLM/fallback classifier improve document type when OCR text gives stronger evidence.
    llm_label = (llm.get("document_type") or doc.get("document_label") or "General Document").strip()
    label_to_key = {v.lower(): k for k, v in DOC_LABELS.items()}
    document_type = label_to_key.get(llm_label.lower(), doc.get("document_type", "general_document"))
    document_label = DOC_LABELS.get(document_type, llm_label or doc.get("document_label", "General Document"))

    llm_actions = llm.get("action_items") or []
    merged_actions = list(dict.fromkeys([str(x) for x in (llm_actions + actions) if str(x).strip()]))[:8]
    merged_dates = list(dict.fromkeys([str(x) for x in ((llm.get("dates") or []) + dates) if str(x).strip()]))[:10]
    merged_amounts = list(dict.fromkeys([str(x) for x in ((llm.get("amounts") or []) + amounts) if str(x).strip()]))[:10]

    risk = attachment_risk(filename, mime_type, sender_band, source_folder, document_type)
    boost = attachment_priority_boost(document_type, risk.get("risk_level", "low"), merged_actions, max(float(doc.get("confidence", 0.35) or 0.35), float(llm.get("confidence", 0.0) or 0.0)))

    return {
        "filename": filename,
        "file_type": file_type,
        "document_type": document_type,
        "document_label": document_label,
        "document_confidence": max(float(doc.get("confidence", 0.35) or 0.35), float(llm.get("confidence", 0.0) or 0.0)),
        "matched_terms": doc.get("matched_terms", []),
        "title": llm.get("title", ""),
        "summary": llm.get("summary") or _summary_from_fields(file_type, doc, extracted, dates, amounts, actions),
        "key_details": llm.get("key_details", []),
        "ids": llm.get("ids", []),
        "business_value": llm.get("business_value", ""),
        "extracted_text": extracted[:5000],
        "analysis_scope": "quick",
        "pages_limited": True,
        "action_items": merged_actions,
        "dates": merged_dates,
        "amounts": merged_amounts,
        "action_required": bool(llm.get("action_required")),
        "priority_boost": boost,
        "priority_reason": llm.get("priority_reason", ""),
        "reply_context": llm.get("reply_context") or f"Attachment analyzed: {filename} is a {document_label}.",
        "llm_summary_used": bool(llm.get("llm_summary_used", False)),
        **risk,
        "safe_to_preview": file_type not in {"archive", "risky_executable"},
    }
